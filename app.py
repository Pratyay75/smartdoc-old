import os
import json
import uuid
import logging
from datetime import timedelta, datetime
from flask import Flask, request, jsonify
from flask import send_from_directory
from flask_cors import CORS
from dotenv import load_dotenv
from pymongo import MongoClient
from werkzeug.security import generate_password_hash, check_password_hash
import fitz  # PyMuPDF
from dateutil import parser as dateparser
from openai import AzureOpenAI
import requests
from Analytics import ( 
    calculate_analytics,
)
from ingest_pdf import push_chunks_to_search, extract_chunks
# ------------------ CONFIG ------------------
load_dotenv() 

app = Flask(__name__, static_folder="frontend/build", static_url_path="")

from flask_cors import CORS

CORS(app, supports_credentials=True, resources={r"/*": {
    "origins": [
        "http://localhost:3000",
        "https://smartdoc.azurewebsites.net"
    ],
    "allow_headers": ["Content-Type", "Authorization"],
    "methods": ["GET", "POST", "OPTIONS", "PUT", "DELETE"]
}})

# -------------- Azure OpenAI & Azure Search Config --------------
EMBEDDING_URL = f"{os.getenv('AZURE_OPENAI_ENDPOINT')}openai/deployments/{os.getenv('AZURE_EMBEDDING_DEPLOYMENT')}/embeddings?api-version={os.getenv('AZURE_API_VERSION')}"
EMBEDDING_HEADERS = {
    "api-key": os.getenv("AZURE_OPENAI_API_KEY"),
    "Content-Type": "application/json"
}

SEARCH_HEADERS = {
    "Content-Type": "application/json",
    "api-key": os.getenv("AZURE_SEARCH_API_KEY")
}


#----------------------------------------

logging.basicConfig(level=logging.INFO)

# Azure OpenAI Setup
client_azure = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_version=os.getenv("AZURE_API_VERSION")
)
DEPLOYMENT_NAME = os.getenv("AZURE_GPT_DEPLOYMENT")
# MongoDB Setup
mongo_client = MongoClient(os.getenv("MONGO_URI"))
db = mongo_client["pdf_data"]
pdf_collection = db["extracted_data"]
users_collection = db["users"]

# Azure Search Setup
AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_ENDPOINT")
AZURE_SEARCH_KEY = os.getenv("AZURE_SEARCH_API_KEY")
AZURE_SEARCH_INDEX = os.getenv("AZURE_SEARCH_INDEX")


# ------------------ Document classification helper ------------------


# ------------------ SIGNUP ------------------
@app.route("/signup", methods=["POST"])
def signup():
    data = request.json
    if users_collection.find_one({"email": data["email"]}):
        return jsonify({"error": "Email already registered"}), 400

    hashed_pw = generate_password_hash(data["password"])
    users_collection.insert_one({
        "name": data.get("name", data["email"].split("@")[0]),
        "email": data["email"],
        "password": hashed_pw
    })
    return jsonify({"message": "Signup successful"})


# ------------------ LOGIN ------------------
@app.route("/login", methods=["POST"])
def login():
    try:
        data = request.get_json(force=True)
        email = data.get("email")
        password = data.get("password")

        user = users_collection.find_one({"email": email})
        if not user or not check_password_hash(user.get("password", ""), password):
            return jsonify({"error": "Invalid credentials"}), 401

        return jsonify({
            "token": str(user["_id"]),
            "name": user.get("name", email.split("@")[0])
        })
    except Exception as e:
        logging.error(f"Login error: {str(e)}")
        return jsonify({"error": "Server error"}), 500


#---------------------------------------------------
def format_ai_data(ai_data):
    try:
        name_conf = ai_data.get("policyholderName_confidence", 0)
        amount_conf = ai_data.get("premiumAmount_confidence", 0)
        date_conf = ai_data.get("issueDate_confidence", 0)

        field_confidences = {
            "name": name_conf,
            "contractAmount": amount_conf,
            "issueDate": date_conf
        }

        total = sum(field_confidences.values())
        count = len(field_confidences)
        accuracy = round(total / count, 2) if count > 0 else 0

        ai_data["field_confidences"] = field_confidences
        ai_data["accuracy"] = accuracy

        return ai_data

    except Exception as e:
        logging.warning(f" format_ai_data() failed: {e}")
        return ai_data

#-------------------------image to text -------------------------
def extract_text_with_azure(file_path):
    from azure.ai.formrecognizer import DocumentAnalysisClient
    from azure.core.credentials import AzureKeyCredential

    endpoint = os.getenv("AZURE_OCR_ENDPOINT")
    key = os.getenv("AZURE_OCR_KEY")

    document_analysis_client = DocumentAnalysisClient(
        endpoint=endpoint, credential=AzureKeyCredential(key)
    )

    try:
        with open(file_path, "rb") as f:
            poller = document_analysis_client.begin_analyze_document(
                "prebuilt-layout",
                document=f
            )
        result = poller.result()

        all_text = []
        page_texts = {}

        # Prefer paragraphs with page numbers
        if hasattr(result, "paragraphs"):
            for para in result.paragraphs:
                if getattr(para, "confidence", 1.0) >= 0.6:
                    all_text.append(para.content)
                    if hasattr(para, "bounding_regions") and para.bounding_regions:
                        page_num = para.bounding_regions[0].page_number
                        page_texts.setdefault(page_num, []).append(para.content)

        # Fallback: entire raw text
        if not all_text and hasattr(result, "content"):
            all_text.append(result.content)
            page_texts[1] = [result.content]

        return {
            "full_text": "\n".join(all_text).strip(),
            "pages": {p: " ".join(txts) for p, txts in page_texts.items()}
        }

    except Exception as e:
        logging.error(f"❌ Azure OCR failed: {e}")
        return {"full_text": "", "pages": {}}

# ------------------ PDF EXTRACTION ------------------
@app.route("/extract", methods=["POST"])
def extract_data():
    try:
        file = request.files.get("pdf")
        if not file:
            return jsonify({"error": "No PDF file provided"}), 400

        pdf_id = str(uuid.uuid4())
        if not file or file.filename == "":
            logging.error("❌ No PDF file uploaded.")
            return jsonify({"error": "No PDF file uploaded"}), 400

        filename = (file.filename or f"uploaded_{uuid.uuid4()}.pdf").replace(" ", "_")

        from azure.storage.blob import BlobServiceClient
        from io import BytesIO

        BLOB_CONN_STR = os.getenv("AZURE_BLOB_CONNECTION_STRING")
        BLOB_CONTAINER = os.getenv("AZURE_STORAGE_CONTAINER")

        blob_service = BlobServiceClient.from_connection_string(BLOB_CONN_STR)
        container_client = blob_service.get_container_client(BLOB_CONTAINER)

        # Upload to Azure Blob
        blob_name = f"{uuid.uuid4()}_{filename}"
        blob_client = container_client.get_blob_client(blob_name)
        blob_client.upload_blob(file, overwrite=True)

        # Download file into memory
        file_stream = BytesIO()
        blob_client.download_blob().readinto(file_stream)
        file_stream.seek(0)

        text = ""
        word_count = 0
        empty_pages = 0
        page_map = {}

        pdf_file = fitz.open(stream=file_stream.read(), filetype="pdf")

        # Push chunks to Azure Cognitive Search
        chunks = [page.get_text().strip() for page in pdf_file if page.get_text().strip()]
        push_chunks_to_search(chunks, source_name=filename)

        page_count = len(pdf_file)

        # Count words & detect empty pages
        file_stream.seek(0)
        pdf_file = fitz.open(stream=file_stream.read(), filetype="pdf")
        for page in pdf_file:
            page_text = page.get_text().strip()
            if not page_text:
                empty_pages += 1
            else:
                text += page_text + "\n"
                word_count += len(page_text.split())

        empty_ratio = empty_pages / page_count
        ocr_used = False

        # OCR fallback
        if word_count < 30 or empty_ratio > 0.5:
            logging.warning(
                f" Detected scanned PDF (word_count={word_count}, empty_pages={empty_pages}/{page_count}) — using Azure OCR fallback."
            )
            import tempfile
            file_stream.seek(0)  # reset before saving
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file_stream.read())
                tmp_path = tmp.name

            ocr_result = extract_text_with_azure(tmp_path)
            text = ocr_result["full_text"]
            page_map = ocr_result["pages"]
            logging.info(" OCR Extracted Text:\n" + text[:1000])
            word_count = len(text.split())
            ocr_used = True

        prompt = f"""
        You are a professional document parser AI. Your task is to extract **structured information** 
        from health insurance policy documents, regardless of how messy or inconsistent the text may be. 
        Use the following schema to return the extracted data as pure JSON only (no extra text):

        {{
            "policyholderName": {{ "value": string | null, "confidence": integer }},
            "issueDateRaw": string | null,
            "issueDate": {{ "value": string | null, "confidence": integer }},
            "expirationDateRaw": string | null,
            "expirationDate": {{ "value": string | null, "confidence": integer }},
            "providerName": {{ "value": string | null, "confidence": integer }},
            "policyholderAddress": {{ "value": string | null, "confidence": integer }},
            "policyNumber": {{ "value": string | null, "confidence": integer }},
            "premiumAmount": {{ "value": string | null, "confidence": integer }},
            "deductibles": {{ "value": string | null, "confidence": integer }},
            "termsAndExclusions": list of strings | null
        }}

        ... Here is the text to extract from:
        {text}
        """

        response = client_azure.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=[
                {"role": "system", "content": "You extract structured data from contracts, even if the format is messy."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2
        )

        extracted_data = response.choices[0].message.content.strip()

        # Clean JSON
        import re
        cleaned = re.sub(r"^(?:```json)?|```$", "", extracted_data.strip(), flags=re.MULTILINE).strip()
        match = re.search(r"\{.*\}", cleaned, re.DOTALL)
        cleaned = match.group(0) if match else extracted_data

        # Parse and flatten JSON
        try:
            parsed_data = json.loads(cleaned)
            flattened = {
                "policyholderName": parsed_data.get("policyholderName", {}).get("value"),
                "policyholderName_confidence": parsed_data.get("policyholderName", {}).get("confidence", 0),
                "issueDateRaw": parsed_data.get("issueDateRaw"),
                "issueDate": parsed_data.get("issueDate", {}).get("value"),
                "issueDate_confidence": parsed_data.get("issueDate", {}).get("confidence", 0),
                "expirationDateRaw": parsed_data.get("expirationDateRaw"),
                "expirationDate": parsed_data.get("expirationDate", {}).get("value"),
                "expirationDate_confidence": parsed_data.get("expirationDate", {}).get("confidence", 0),
                "providerName": parsed_data.get("providerName", {}).get("value"),
                "providerName_confidence": parsed_data.get("providerName", {}).get("confidence", 0),
                "policyholderAddress": parsed_data.get("policyholderAddress", {}).get("value"),
                "policyholderAddress_confidence": parsed_data.get("policyholderAddress", {}).get("confidence", 0),
                "policyNumber": parsed_data.get("policyNumber", {}).get("value"),
                "policyNumber_confidence": parsed_data.get("policyNumber", {}).get("confidence", 0),
                "premiumAmount": parsed_data.get("premiumAmount", {}).get("value"),
                "premiumAmount_confidence": parsed_data.get("premiumAmount", {}).get("confidence", 0),
                "deductibles": parsed_data.get("deductibles", {}).get("value"),
                "deductibles_confidence": parsed_data.get("deductibles", {}).get("confidence", 0),
                "termsAndExclusions": parsed_data.get("termsAndExclusions"),
            }

            # Fallback regex logic
            if not flattened.get("premiumAmount"):
                match = re.search(
                    r"(sum assured|total benefit|maturity amount)[^\n]*?(Rs\.?\s*[\d,]+)",
                    text,
                    re.IGNORECASE
                )
                if match:
                    flattened["premiumAmount"] = match.group(2).strip()
                    flattened["premiumAmount_confidence"] = 75

            if not flattened.get("deductibles"):
                match = re.search(
                    r"(premium(?: per| payable)?(?:.*)?)[^\n]*?(Rs\.?\s*[\d,]+\s*(?:monthly|quarterly|annually|yearly)?)",
                    text,
                    re.IGNORECASE
                )
                if match:
                    flattened["deductibles"] = match.group(2).strip()
                    flattened["deductibles_confidence"] = 70

            # Format extracted dates to DD-MM-YYYY
            for field in ["issueDate", "expirationDate"]:
                if flattened.get(field):
                    try:
                        dt = dateparser.parse(flattened[field], fuzzy=True)
                        flattened[field] = dt.strftime("%d-%m-%Y")
                    except Exception as e:
                        logging.warning(f"⚠️ Could not format {field}: {e}")

            parsed_data = format_ai_data(flattened)

            # ---------- NEW: field → page mapping ----------
            field_page_map = {}

            def find_page_for_value(val):
                if not val or not page_map:
                    return None
                val_lower = val.lower()
                for p, ptext in page_map.items():
                    if val_lower in ptext.lower():
                        return p
                return None

            for key in ["policyholderName", "providerName", "policyNumber", "premiumAmount", "deductibles", "policyholderAddress"]:
                field_page_map[f"{key}_page"] = find_page_for_value(flattened.get(key))

            field_page_map["issueDate_page"] = find_page_for_value(flattened.get("issueDateRaw"))
            field_page_map["expirationDate_page"] = find_page_for_value(flattened.get("expirationDateRaw"))

        except json.JSONDecodeError:
            logging.error(f"⚠️ Invalid JSON from model: {cleaned}")
            parsed_data = {"raw_output": extracted_data}
            field_page_map = {}

        # Save to MongoDB
        user_id = request.form.get("user_id")
        if not user_id:
            return jsonify({"error": "Missing user_id in form data"}), 400

        pdf_collection.insert_one({
            "pdf_id": pdf_id,
            "pdfName": filename,
            "ai_data": parsed_data,
            "pageCount": page_count,
            "wordCount": word_count,
            "timestamp": datetime.utcnow(),
            "user_id": user_id
        })

        return jsonify({
            "pdf_id": pdf_id,
            "ocr_used": ocr_used,
            **parsed_data,
            **field_page_map
        })

    except Exception as e:
        logging.error(f"❌ Error during extraction: {str(e)}")
        return jsonify({"error": str(e)}), 500

# ------------------ SAVE EDITED DATA ------------------
@app.route("/save", methods=["POST"])
def save():
    data = request.get_json()
    user_id = data.get("user_id")
    pdf_id = data.get("pdf_id")
    updated_fields = data.get("user_updated_data")

    if not pdf_id or not updated_fields:
        return jsonify({"error": "Missing pdf_id or updated data"}), 400

    # Format issueDate if present
    if "issueDate" in updated_fields:
        try:
            dt = dateparser.parse(updated_fields["issueDate"], fuzzy=True)
            updated_fields["issueDate"] = dt.strftime("%d-%m-%Y")
        except Exception as e:
            logging.warning(f"Could not parse issueDate in save(): {e}")

    # Get the original document to compare
    existing = pdf_collection.find_one({"pdf_id": pdf_id})
    if not existing:
        return jsonify({"error": "PDF not found"}), 404

    ai_data = existing.get("ai_data", {})
    # Compare only changed fields
    changes = {k: v for k, v in updated_fields.items() if ai_data.get(k) != v}

    if not changes:
        return jsonify({"message": "Data Saved"}), 200

    result = pdf_collection.update_one(
    {"pdf_id": pdf_id},
    {
        "$set": {
            "user_updated_data": changes,
            "user_id": user_id,  # ensure update preserves user
            "timestamp": datetime.utcnow()
        }
    }
)


    return jsonify({"message": "User updated data saved successfully"})




# ------------------ CHATBOT ------------------
def query_azure_search(question, top_k=5):
    url = f"{AZURE_SEARCH_ENDPOINT}/indexes/{AZURE_SEARCH_INDEX}/docs/search?api-version=2023-07-01-Preview"
    headers = {"Content-Type": "application/json", "api-key": AZURE_SEARCH_KEY}
    body = {"search": question, "top": top_k}
    try:
        response = requests.post(url, headers=headers, json=body)
        response.raise_for_status()
        results = response.json()
        return [doc["content"] for doc in results.get("value", [])]
    except Exception as e:
        print("❌ Azure Search Query Failed:", e)
        return []


@app.route("/chat", methods=["POST"])
def chat():
    data = request.json
    pdf_id = data.get("pdf_id")
    question = data.get("question")

    record = pdf_collection.find_one({"pdf_id": pdf_id})
    if not record:
        return jsonify({"error": "PDF data not found"}), 404

    ai_summary = json.dumps(record.get("ai_data", {}), indent=2)
    search_chunks = query_azure_search(question)
    full_text = "\n\n---\n\n".join(search_chunks) if search_chunks else ai_summary

    prompt = f"""
You are  —Chatbot a smart, human-like assistant trained to help users understand complex PDFs such as contracts, insurance policies, business reports, or legal documents.

🎯 Your Goal:
Help the user by answering their question **only using the content of the provided PDF**. Be friendly, clear, and act like a real assistant — not a machine.

---

🧠 Behavior Rules:
- Be professional, conversational, and accurate.
- Use ONLY the content in the PDF to answer.
- If something is not clearly mentioned, say so politely.
- Do not assume or guess beyond what’s written.

📌 Formatting Rules:
- If the user asks for **bullet points, lists, dates, exclusions, or summary points**, format them as:
  - Each item starts with a dash (-).
  - Each item is on a new line.
  - Leave a blank line between items for better readability.
- If the user asks for **steps or instructions**, format them with:
  1. Numbered steps
  2. Clear spacing
  3. Proper punctuation
- If the user asks for a **specific value** (e.g., date, name, amount):
  → Give a short, direct, clear sentence.
- Do NOT return any code, JSON, or technical symbols.

---

📄 PDF Content:
{full_text}

❓ User’s Question:
{question}

---

💬 Your Answer:
(Reply naturally like a helpful assistant would. Avoid sounding robotic.)
"""


    try:
        response = client_azure.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=[
                {"role": "system", "content": "You are a conversational assistant answering based on PDF content."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5
        )
        answer = response.choices[0].message.content.strip()
        return jsonify({"answer": answer})
    except Exception as e:
        logging.error(f"❌ Error in chatbot: {str(e)}")
        return jsonify({"error": str(e)}), 500




#-------------------analytics---------------
@app.route("/analytics", methods=["POST"])
def get_user_analytics():
    try:
        data = request.get_json()
        user_id = data.get("user_id")
        period = data.get("filter", "month")  

        if not user_id:
            return jsonify({"error": "Missing user_id"}), 400

        from Analytics import calculate_analytics
        analytics_data = calculate_analytics(pdf_collection, period=period, user_id=user_id)
        return jsonify(analytics_data)

    except Exception as e:
        logging.error(f"❌ Analytics route error: {str(e)}")
        return jsonify({"error": "Failed to calculate analytics"}), 500


@app.route("/analytics/trends", methods=["POST"])
def analytics_trends():
    data = request.get_json()
    user_id = data.get("user_id")
    filter_by = data.get("filter", "month")

    if not user_id:
        return jsonify({"error": "Missing user ID"}), 400

    now = datetime.now()
    if filter_by == "day":
        start_time = now.replace(hour=0, minute=0, second=0, microsecond=0)
    elif filter_by == "week":
        start_time = now - timedelta(days=7)
    elif filter_by == "month":
        start_time = now - timedelta(days=30)
    else:
        start_time = datetime.min

    pipeline = [
        {
            "$match": {
                "user_id": user_id,
                "timestamp": {"$gte": start_time}
            }
        },
        {
            "$group": {
                "_id": {
                    "year": {"$year": "$timestamp"},
                    "month": {"$month": "$timestamp"},
                    "day": {"$dayOfMonth": "$timestamp"},
                },
                "avg_accuracy": {"$avg": "$ai_data.accuracy"}
            }
        },
        {
            "$sort": {"_id": 1}
        }
    ]

    results = list(pdf_collection.aggregate(pipeline))

    trend = []
    for r in results:
        y, m, d = r["_id"]["year"], r["_id"]["month"], r["_id"]["day"]
        date_str = f"{d:02d}-{m:02d}-{y}"
        trend.append({"date": date_str, "avg_accuracy": r["avg_accuracy"]})

    return jsonify({"trend": trend})

@app.route("/analytics/pdf-details", methods=["POST"])
def analytics_pdf_details():
    data = request.get_json()
    user_id = data.get("user_id")

    if not user_id:
        return jsonify({"error": "Missing user ID"}), 400

    pdfs = list(pdf_collection.find(
    {"user_id": user_id},
    {"pdfName": 1, "ai_data": 1, "timestamp": 1, "pageCount": 1, "wordCount": 1}
).sort("timestamp", -1))


    for pdf in pdfs:
        pdf["_id"] = str(pdf["_id"])

        # Extract from ai_data if not top-level
        ai_data = pdf.get("ai_data", {})
        pdf["accuracy"] = ai_data.get("accuracy")
        pdf["field_confidences"] = ai_data.get("field_confidences", {})

        # Format timestamp for display
        if "timestamp" in pdf:
            pdf["timestamp"] = pdf["timestamp"].strftime("%d-%m-%Y %H:%M")
    return jsonify({"pdfs": pdfs})
# --------------------------- PDF COMPARE ---------------------------
import fitz  # PyMuPDF
import re
from dateutil import parser as dateparser
from flask import Flask, request, jsonify
from rapidfuzz import fuzz
import diff_match_patch as dmp_module

# ---------- Config ----------
PARA_MATCH_THRESHOLD = 80   # lower to catch minor changes
LINE_MATCH_THRESHOLD = 85   # slightly lower for OCR tolerance
MIN_PARTIAL_THRESHOLD = 60

DATE_RE = re.compile(r"\b(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})\b")
NUMERIC_RE = re.compile(r"\b\d+(?:\.\d+)?\b")  # matches integers and decimals

# ---------- Helpers ----------

def normalize_whitespace(text: str) -> str:
    text = text.replace('\r', '\n')
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(ln.strip() for ln in text.splitlines())
    return text.strip()

def normalize_dates(text: str) -> str:
    def repl(m):
        try:
            d = dateparser.parse(m.group(0), dayfirst=True)
            return d.strftime("%d-%m-%Y")
        except Exception:
            return m.group(0)
    return DATE_RE.sub(repl, text)

def extract_paragraphs_from_pdf_bytes(file_bytes: bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    paras = []
    for page in doc:
        raw = page.get_text("text") or ""
        raw = normalize_whitespace(raw)
        # Split on blank lines OR newline + capital letter
        for p in re.split(r"(?:\n\s*\n)|(?:\n(?=[A-Z]))", raw):
            p = p.strip()
            if p:
                p = normalize_dates(p)
                paras.append(p)
    doc.close()
    return paras

def numbers_changed(a: str, b: str) -> bool:
    """Check if numeric values differ between strings."""
    nums_a = re.findall(r"\b\d+(?:\.\d+)?\b", a)
    nums_b = re.findall(r"\b\d+(?:\.\d+)?\b", b)
    return nums_a != nums_b

def dates_changed(a: str, b: str) -> bool:
    """Check if date values differ between strings (after normalization)."""
    dates_a = [normalize_dates(m) for m in DATE_RE.findall(a)]
    dates_b = [normalize_dates(m) for m in DATE_RE.findall(b)]
    return dates_a != dates_b

def word_level_diff_html(a: str, b: str) -> str:
    # NEW: if numeric or date values differ, mark whole thing removed/new
    if numbers_changed(a, b) or dates_changed(a, b):
        return f'<span class="removed">{a}</span><span class="new">{b}</span>'
    
    dmp = dmp_module.diff_match_patch()
    diffs = dmp.diff_main(a, b)
    dmp.diff_cleanupSemantic(diffs)
    parts = []
    for op, data in diffs:
        txt = data.replace("\n", "<br/>")
        if op == 0:
            parts.append(f'<span class="same">{txt}</span>')
        elif op == -1:
            parts.append(f'<span class="removed">{txt}</span>')
        elif op == 1:
            parts.append(f'<span class="new">{txt}</span>')
    return "".join(parts)

# ---------- Main Compare Route ----------

@app.route("/compare", methods=["POST"])
def compare_pdfs():
    try:
        pdf1 = request.files.get("pdf1")
        pdf2 = request.files.get("pdf2")
        if not pdf1 or not pdf2:
            return jsonify({"error": "Both pdf1 and pdf2 are required"}), 400

        paras1 = extract_paragraphs_from_pdf_bytes(pdf1.read())
        paras2 = extract_paragraphs_from_pdf_bytes(pdf2.read())

        matched_2 = set()
        html_blocks = []

        # Match paragraphs regardless of order
        for p1 in paras1:
            best_score = -1
            best_j = None
            for j, p2 in enumerate(paras2):
                if j in matched_2:
                    continue
                score = fuzz.ratio(p1, p2)
                if score > best_score:
                    best_score = score
                    best_j = j

            if best_score >= PARA_MATCH_THRESHOLD and best_j is not None:
                wl_html = word_level_diff_html(p1, paras2[best_j])
                html_blocks.append(f'<div class="para same">{wl_html}</div>')
                matched_2.add(best_j)
            else:
                best_para_j = None
                best_para_score = -1
                for j, p2 in enumerate(paras2):
                    if j in matched_2:
                        continue
                    score = fuzz.partial_ratio(p1, p2)
                    if score > best_para_score:
                        best_para_score = score
                        best_para_j = j

                if best_para_j is None or best_para_score < MIN_PARTIAL_THRESHOLD:
                    html_blocks.append(f'<div class="para removed">{p1}</div>')
                else:
                    p2 = paras2[best_para_j]
                    matched_2.add(best_para_j)

                    lines1 = [ln.strip() for ln in p1.splitlines() if ln.strip()]
                    lines2 = [ln.strip() for ln in p2.splitlines() if ln.strip()]
                    matched_lines2 = set()
                    para_html = ['<div class="para">']

                    for l1 in lines1:
                        best_lscore = -1
                        best_lidx = None
                        for idx2, l2 in enumerate(lines2):
                            if idx2 in matched_lines2:
                                continue
                            score = fuzz.ratio(l1, l2)
                            if score > best_lscore:
                                best_lscore = score
                                best_lidx = idx2
                        if best_lscore >= LINE_MATCH_THRESHOLD and best_lidx is not None:
                            wl_html = word_level_diff_html(l1, lines2[best_lidx])
                            para_html.append(f'<div class="line same">{wl_html}</div>')
                            matched_lines2.add(best_lidx)
                        elif best_lscore >= MIN_PARTIAL_THRESHOLD and best_lidx is not None:
                            wl_html = word_level_diff_html(l1, lines2[best_lidx])
                            para_html.append(f'<div class="line partial">{wl_html}</div>')
                            matched_lines2.add(best_lidx)
                        else:
                            para_html.append(f'<div class="line removed">{l1}</div>')

                    # New lines in p2 not matched
                    for idx2, l2 in enumerate(lines2):
                        if idx2 not in matched_lines2:
                            para_html.append(f'<div class="line new">{l2}</div>')

                    para_html.append('</div>')
                    html_blocks.append("".join(para_html))

        # Any paras in PDF2 not matched at all → NEW
        for j, p2 in enumerate(paras2):
            if j not in matched_2:
                html_blocks.append(f'<div class="para new">{p2}</div>')

        final_html = '<div class="compare-output">' + "\n".join(html_blocks) + '</div>'
        return jsonify({"html_result": final_html}), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ------------------ MULTI-DOC UPLOAD ------------------
from azure.storage.blob import BlobServiceClient
import tempfile

@app.route("/upload-multi-doc", methods=["POST"])
def upload_multi_doc():
    try:
        files = request.files.getlist("files")
        if not files:
            return jsonify({"error": "No files provided"}), 400

        BLOB_CONN_STR = os.getenv("AZURE_BLOB_CONNECTION_STRING")
        BLOB_CONTAINER = os.getenv("AZURE_STORAGE_CONTAINER")
        blob_service = BlobServiceClient.from_connection_string(BLOB_CONN_STR)
        container_client = blob_service.get_container_client(BLOB_CONTAINER)

        uploaded_docs = []

        for file in files:
            filename = file.filename
            blob_name = f"{uuid.uuid4()}_{filename.replace(' ', '_')}"
            blob_client = container_client.get_blob_client(blob_name)
            blob_client.upload_blob(file, overwrite=True)
            logging.info(f"✅ Uploaded to Azure Blob: {blob_name}")

            # Reset file pointer for text extraction
            file.stream.seek(0)
            file_bytes = file.read()

            status = "Uploaded"
            try:
                from ingest_multi_doc import process_blob
                process_blob(blob_name, filename)  # pass filename for better logging
            except Exception as ingest_err:
                logging.error(f"❌ Failed to process {filename}: {str(ingest_err)}")
                status = "Failed"

            uploaded_docs.append({
                "name": filename,
                "date": datetime.utcnow().strftime("%d-%m-%Y %H:%M"),
                "status": status,
                "size": f"{len(file_bytes) / 1024:.1f} KB",
                "blob_name": blob_name
            })

        return jsonify({"documents": uploaded_docs})

    except Exception as e:
        logging.error(f"❌ Multi-doc upload error: {str(e)}")
        return jsonify({"error": str(e)}), 500


from flask import request, jsonify
from azure.storage.blob import BlobServiceClient
import os, requests, logging

SEARCH_HEADERS = {
    "Content-Type": "application/json",
    "api-key": os.getenv("AZURE_SEARCH_API_KEY")
}

def get_doc_ids_by_blob(blob_name):
    """
    Query Azure Search to get all document IDs for a given blob name.
    Escapes single quotes for OData.
    """
    safe_blob = blob_name.replace("'", "''")
    search_url = f"{os.getenv('AZURE_SEARCH_ENDPOINT')}/indexes/{os.getenv('AZURE_MULTI_DOC_INDEX')}/docs/search?api-version=2023-07-01-Preview"
    query_payload = {
        "filter": f"metadata eq 'source:{safe_blob}'",
        "select": "id",
        "top": 1000
    }
    res = requests.post(search_url, headers=SEARCH_HEADERS, json=query_payload)
    if res.status_code == 200:
        docs = res.json().get("value", [])
        return [doc["id"] for doc in docs]
    else:
        logging.error(f"❌ Failed to get doc IDs from Azure Search for {blob_name}: {res.text}")
        return []

def delete_from_search(blob_name):
    """Delete all indexed chunks in Azure Search for a given blob."""
    doc_ids = get_doc_ids_by_blob(blob_name)
    if not doc_ids:
        logging.warning(f"⚠ No matching chunks found in Azure Search for blob {blob_name}")
        return

    delete_url = f"{os.getenv('AZURE_SEARCH_ENDPOINT')}/indexes/{os.getenv('AZURE_MULTI_DOC_INDEX')}/docs/index?api-version=2023-07-01-Preview"
    delete_payload = {
        "value": [{"@search.action": "delete", "id": doc_id} for doc_id in doc_ids]
    }
    res = requests.post(delete_url, headers=SEARCH_HEADERS, json=delete_payload)
    logging.info(f"🔄 Azure Search delete response for {blob_name}: {res.status_code} - {res.text}")


@app.route("/delete-blob", methods=["POST"])
def delete_blob():
    try:
        data = request.get_json()
        blob_name = data.get("blob_name")
        if not blob_name:
            return jsonify({"error": "Missing blob_name"}), 400

        # ---------- 1. Delete from Azure Blob ----------
        BLOB_CONN_STR = os.getenv("AZURE_BLOB_CONNECTION_STRING")
        BLOB_CONTAINER = os.getenv("AZURE_STORAGE_CONTAINER")
        blob_service = BlobServiceClient.from_connection_string(BLOB_CONN_STR)
        blob_client = blob_service.get_blob_client(BLOB_CONTAINER, blob_name)
        blob_client.delete_blob()
        logging.info(f"✅ Deleted blob: {blob_name}")

        # ---------- 2. Delete from Azure Search ----------
        delete_from_search(blob_name)

        return jsonify({"message": f"Blob '{blob_name}' and its indexed chunks deleted successfully"})

    except Exception as e:
        logging.error(f"❌ Delete blob error: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route("/delete-multiple-blobs", methods=["POST"])
def delete_multiple_blobs():
    try:
        data = request.get_json()
        blob_names = data.get("blob_names", [])
        if not blob_names:
            return jsonify({"error": "Missing blob_names list"}), 400

        BLOB_CONN_STR = os.getenv("AZURE_BLOB_CONNECTION_STRING")
        BLOB_CONTAINER = os.getenv("AZURE_STORAGE_CONTAINER")
        blob_service = BlobServiceClient.from_connection_string(BLOB_CONN_STR)

        deleted_count = 0
        for blob_name in blob_names:
            try:
                blob_client = blob_service.get_blob_client(BLOB_CONTAINER, blob_name)
                blob_client.delete_blob()
                logging.info(f"✅ Deleted blob: {blob_name}")

                delete_from_search(blob_name)
                deleted_count += 1
            except Exception as inner_err:
                logging.error(f"❌ Failed to delete blob '{blob_name}': {str(inner_err)}")

        return jsonify({"message": f"Deleted {deleted_count} blobs and cleaned up Azure Search"})

    except Exception as e:
        logging.error(f"❌ Delete multiple blobs error: {str(e)}")
        return jsonify({"error": str(e)}), 500


#----------------new chatbot----------------------
import re

def _tokenize(text):
    return re.findall(r"[a-zA-Z0-9]+", (text or "").lower())

def _score_chunk(chunk_text, q_tokens):
    # simple term-frequency score
    if not chunk_text:
        return 0
    text_tokens = _tokenize(chunk_text)
    if not text_tokens:
        return 0
    hits = sum(text_tokens.count(t) for t in q_tokens if t)
    return hits

def select_top_chunks(docs, question, per_file_cap=8, total_cap=24):
    """
    docs: list of {"filename", "blob_name", "chunks": [ { "content": "..."} or "..." ]}
    Returns a list of top chunks across selected docs.
    """
    q_tokens = _tokenize(question)
    scored = []

    for d in docs:
        filename = d.get("filename") or d.get("name") or "document"
        blob = d.get("blob_name")
        raw_chunks = d.get("chunks", [])
        # support both formats: list[str] or list[{"content": "..."}]
        norm = []
        for c in raw_chunks:
            if isinstance(c, dict):
                norm.append(c.get("content", ""))
            else:
                norm.append(str(c))

        # per-file ranking
        file_scored = []
        for ch in norm:
            s = _score_chunk(ch, q_tokens)
            # slight boost if file name words appear
            s += _score_chunk(filename, q_tokens) * 0.2
            file_scored.append((s, filename, blob, ch))

        # keep best N per file
        file_scored.sort(key=lambda x: x[0], reverse=True)
        scored.extend(file_scored[:per_file_cap])

    # global cap
    scored.sort(key=lambda x: x[0], reverse=True)
    top = [dict(score=s, filename=f, blob_name=b, content=c) for s, f, b, c in scored[:total_cap]]
    return top

# ------------------ CHATBOT (Multi-Doc restricted) ------------------

@app.route("/chat-multidoc", methods=["POST"])
def chat_multidoc():
    try:
        data = request.get_json(force=True)
        blob_names = data.get("blob_names", [])
        question = data.get("question", "")

        if not blob_names:
            return jsonify({"error": "No blob_names provided"}), 400
        if not question.strip():
            return jsonify({"error": "Question is empty"}), 400

        # 1️⃣ Get embedding of the question
        embed_payload = {"input": question, "model": os.getenv("AZURE_EMBEDDING_DEPLOYMENT")}
        embed_res = requests.post(EMBEDDING_URL, headers=EMBEDDING_HEADERS, json=embed_payload)
        if embed_res.status_code != 200:
            return jsonify({"error": "Embedding failed"}), 500
        question_vector = embed_res.json()["data"][0]["embedding"]

        # Escape blob_names for OData filter
        safe_blob_names = [bn.replace("'", "''") for bn in blob_names]
        filter_str = " or ".join([f"metadata eq 'source:{bn}'" for bn in safe_blob_names])

        # 2️⃣ Query Azure Cognitive Search (Hybrid: keyword + vector)
        search_payload = {
            "search": question,
            "vector": {
                "value": question_vector,
                "fields": "embedding",
                "k": 20
            },
            "filter": filter_str,
            "select": "content,metadata,filename"
        }
        search_url = f"{os.getenv('AZURE_SEARCH_ENDPOINT')}/indexes/{os.getenv('AZURE_MULTI_DOC_INDEX')}/docs/search?api-version=2023-07-01-Preview"
        search_res = requests.post(search_url, headers=SEARCH_HEADERS, json=search_payload)
        if search_res.status_code != 200:
            return jsonify({"error": "Search failed"}), 500
        hits = search_res.json().get("value", [])

        if not hits:
            return jsonify({"answer": "No relevant content found in the selected documents."})

        # 3️⃣ Use top chunks directly (like single-PDF flow)
        top_chunks = hits[:8]  # trust Cognitive Search ranking
        context_text = "\n\n".join([h["content"] for h in top_chunks if "content" in h])

        prompt = f"""
You are —Chatbot a smart, human-like assistant trained to help users understand complex PDFs such as contracts, insurance policies, business reports, or legal documents.

🎯 Your Goal:
Help the user by answering their question **only using the content of the provided PDFs**. Be friendly, clear, and act like a real assistant — not a machine.

---

🧠 Behavior Rules:
- Be professional, conversational, and accurate.
- Use ONLY the content in the PDFs to answer.
- If something is not clearly mentioned, say so politely.
- If there is a table, use key-value pairs to answer.
- Do not assume or guess beyond what’s written.

Context:
{context_text}

Question:
{question}

Answer:
"""

        # 4️⃣ Ask Azure OpenAI
        response = client_azure.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=[
                {"role": "system", "content": "Answer based only on the provided context."},
                {"role": "user", "content": prompt}
            ],
            temperature=0
        )
        answer = response.choices[0].message.content.strip()

        return jsonify({"answer": answer})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

#-----------------------------------------------------------------
import os, re, chardet, fitz, smtplib, io, logging
from email.mime.text import MIMEText
from flask import request, jsonify, g
from pymongo import MongoClient
from datetime import datetime
from bson import ObjectId

try:
    import docx
except ImportError:
    docx = None

logging.basicConfig(level=logging.INFO)

# ---------- MongoDB Setup ----------
MONGO_URI = os.getenv("MONGO_URI")
client = MongoClient(MONGO_URI)

db = client["pdf_data"]
categories_col = db["categories"]
users_collection = db["users"]



# ---------- Auth helper ----------
@app.before_request
def load_current_user():
    if request.endpoint in ("login", "signup", "index", "static"):
        return

    auth_header = request.headers.get("Authorization")
    g.user_id = None
    if not auth_header or not auth_header.startswith("Bearer "):
        return

    token = auth_header.split(" ")[1].strip()
    try:
        user = users_collection.find_one({"_id": ObjectId(token)})
        if user:
            g.user_id = user["email"]
    except Exception as e:
        logging.error(f"Auth error: {e}")

def _safe_lower(s: str) -> str:
    return (s or "").lower()


EMAIL_REGEX = re.compile(r"(?i)^[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}$")
OTHER = "Other"

# ---------- Extract text ----------
def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = []
            for page in doc:
                txt = page.get_text() or ""
                if txt.strip():
                    text.append(txt)
            doc.close()
            return _safe_lower("\n".join(text))
        except Exception as e:
            logging.error(f"PDF read failed: {e}")
    if name.endswith(".docx") and docx:
        try:
            d = docx.Document(io.BytesIO(file_bytes))
            return _safe_lower(" ".join(p.text for p in d.paragraphs))
        except Exception as e:
            logging.error(f"DOCX read failed: {e}")
    try:
        det = chardet.detect(file_bytes or b"")
        enc = det.get("encoding") or "utf-8"
        return _safe_lower(file_bytes.decode(enc, errors="ignore"))
    except Exception:
        return _safe_lower(file_bytes.decode("utf-8", errors="ignore"))

# ---------- AI Intent ----------
try:
    from openai import AzureOpenAI
    _openai_client = AzureOpenAI(
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        api_version=os.getenv("AZURE_API_VERSION", "2024-12-01-preview"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
    )
except Exception:
    _openai_client = None

def extract_intent(text: str) -> str:
    snippet = (text or "")[:1500]
    prompt = f"""
You are an assistant that analyzes documents.
Summarize the INTENT of this document in one clear sentence.
Document:
{snippet}
""".strip()
    if _openai_client is None:
        return "General inquiry regarding the document content."
    try:
        deployment = os.getenv("AZURE_GPT_DEPLOYMENT", "pdf-gpt")  # Azure deployment name
        resp = _openai_client.chat.completions.create(
            model=deployment,
            messages=[
                {"role": "system", "content": "Return exactly one concise sentence that captures the document's submission intent."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=60,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        logging.error(f"Intent AI error: {e}")
        return "General inquiry regarding the document content."


# ---------- Classification ----------
def classify_document(text: str, user_id: str):
    # Fetch categories for this user
    doc = categories_col.find_one({"user_id": user_id}, {"categories": 1})
    if not doc or not doc.get("categories"):
        return OTHER, None

    scores = {}
    for cat in doc["categories"]:
        cname = (cat.get("name") or "").strip()
        if not cname:
            continue
        scores[cname] = 0
        for kw in cat.get("keywords", []):
            kw = (kw or "").lower().strip()
            if not kw:
                continue
            pattern = r"(?<!\w)" + re.escape(kw) + r"(?!\w)"
            scores[cname] += len(re.findall(pattern, text))

    if not scores:
        return OTHER, None

    best_cat = max(scores, key=scores.get)
    if scores[best_cat] == 0:
        return OTHER, None

    for cat in doc["categories"]:
        if (cat.get("name") or "").strip() == best_cat:
            return best_cat, cat.get("receiver_email")

    return OTHER, None


# ---------- Email ----------
def send_email(receiver, subject, body):
    sender = "yashshrivastava5252@gmail.com"
    try:
        msg = MIMEText(body)
        msg["Subject"] = subject
        msg["From"] = sender
        msg["To"] = receiver
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender, os.getenv("EMAIL_PASSWORD"))
            server.send_message(msg)
        logging.info(f" Email sent to {receiver}")
        return True, None
    except Exception as e:
        logging.error(f"Email send failed: {e}")
        return False, str(e)

# ---------- Routes ----------
@app.route("/classify-docs", methods=["POST"])
def classify_docs():
    if not g.user_id:
        return jsonify({"error": "Unauthorized"}), 401

    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No files provided"}), 400

    results = []
    for fs in files:
        filename = fs.filename or "document"
        file_bytes = fs.read() or b""
        text = extract_text_from_bytes(file_bytes, filename)

        category, _receiver = classify_document(text, g.user_id)
        # Always ensure a value
        category = category or OTHER

        intent = extract_intent(text)
        results.append({
            "name": filename,
            "status": "Done",
            "category": category,
            "intent": intent
        })

    return jsonify({"results": results})


@app.route("/send-classification", methods=["POST"])
def send_classification():
    if not g.user_id:
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json(force=True)
    name = data.get("name") or "Document"
    category = data.get("category") or "Unclassified"
    intent = data.get("intent") or ""
    to_email = (data.get("to_email") or "").strip()
    if not EMAIL_REGEX.match(to_email):
        return jsonify({"error": "Valid recipient email required"}), 400

    subject = f"Document classified as {category}"
    body = f"Document '{name}' was classified as '{category}'.\n\nIntent: {intent}\n"
    ok, err = send_email(to_email, subject, body)
    if not ok:
        return jsonify({"error": f"Email send failed: {err}"}), 500
    return jsonify({"message": "Email sent"})

@app.route("/update-categories", methods=["POST"])
def update_categories():
    if not g.user_id:
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json(force=True)
    category_raw = (data.get("category") or "").strip()
    keywords = data.get("keywords", [])
    receiver_email_raw = (data.get("receiver_email") or "").strip()

    if not category_raw or not keywords or not receiver_email_raw:
        return jsonify({"error": "All fields required"}), 400

    if not EMAIL_REGEX.match(receiver_email_raw):
        return jsonify({"error": "Enter a valid receiver email"}), 400

    name_lc = category_raw.lower()

    # Check duplicate (case-insensitive)
    doc = categories_col.find_one({"user_id": g.user_id}, {"categories": 1})
    if doc and any((c.get("name_lc") or (c.get("name") or "").lower()) == name_lc
                   for c in doc.get("categories", [])):
        return jsonify({"error": "Category name must be unique"}), 409

    # Normalize keywords (trim empty)
    keywords = [str(k).strip() for k in keywords if str(k).strip()]

    categories_col.update_one(
        {"user_id": g.user_id},
        {
            "$setOnInsert": {"user_id": g.user_id},
            "$push": {"categories": {
                "name": category_raw,
                "name_lc": name_lc,
                "keywords": keywords,
                "receiver_email": receiver_email_raw
            }}
        },
        upsert=True
    )
    return jsonify({"message": "Category saved"})


@app.route("/edit-category", methods=["PUT"])
def edit_category():
    if not g.user_id:
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json(force=True)
    name = (data.get("name") or "").strip()
    update = data.get("update", {}) or {}

    set_payload = {}

    # Validate receiver_email if present
    if "receiver_email" in update:
        new_recv = (update["receiver_email"] or "").strip()
        if not EMAIL_REGEX.match(new_recv):
            return jsonify({"error": "Enter a valid receiver email"}), 400
        set_payload["categories.$.receiver_email"] = new_recv

    # Validate & dedupe name if present
    if "name" in update:
        new_name = (update["name"] or "").strip()
        if not new_name:
            return jsonify({"error": "Category name cannot be empty"}), 400
        new_lc = new_name.lower()

        doc = categories_col.find_one({"user_id": g.user_id}, {"categories": 1})
        if doc:
            for c in doc.get("categories", []):
                existing_lc = (c.get("name_lc") or (c.get("name") or "").lower())
                # if another category already has this name
                if existing_lc == new_lc and existing_lc != name.lower():
                    return jsonify({"error": "Category name must be unique"}), 409

        set_payload["categories.$.name"] = new_name
        set_payload["categories.$.name_lc"] = new_lc

    # Update keywords if present
    if "keywords" in update:
        kws = [str(k).strip() for k in (update["keywords"] or []) if str(k).strip()]
        set_payload["categories.$.keywords"] = kws

    if not set_payload:
        return jsonify({"error": "Nothing to update"}), 400

    result = categories_col.update_one(
        {"user_id": g.user_id, "categories.name": name},
        {"$set": set_payload}
    )
    if result.matched_count == 0:
        return jsonify({"error": "Category not found"}), 404

    return jsonify({"message": "Category updated"})


@app.route("/get-categories", methods=["GET"])
def get_categories():
    if not g.user_id:
        return jsonify({"error": "Unauthorized"}), 401
    doc = categories_col.find_one({"user_id": g.user_id}, {"_id": 0, "categories": 1})
    return jsonify({"categories": doc.get("categories", []) if doc else []})

@app.route("/delete-category", methods=["DELETE"])
def delete_category():
    if not g.user_id:
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json(force=True)
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Category name required"}), 400

    # Remove category from user's categories array
    result = categories_col.update_one(
        {"user_id": g.user_id},
        {"$pull": {"categories": {"name": name}}}
    )

    if result.modified_count == 0:
        return jsonify({"error": "Category not found"}), 404

    return jsonify({"message": f"Category '{name}' deleted successfully"})

from flask import Flask, request, jsonify
import fitz  # PyMuPDF
import io


def highlight_mismatches(doc, mismatches, pdf_id):
    """Add highlights to the given doc based on mismatches list."""
    for m in mismatches:
        if pdf_id not in m:
            continue
        w = m[pdf_id]
        page = doc[w["page"] - 1]
        rect = fitz.Rect(w["bbox"])
        highlight = page.add_highlight_annot(rect)
        highlight.set_colors(stroke=(1, 0.6, 0.6) if pdf_id == "pdf1" else (1, 1, 0))
        highlight.update()
    return doc

@app.route("/compare-sbs", methods=["POST"])
def compare_pdfs_sbs():
    if "file1" not in request.files or "file2" not in request.files:
        return jsonify({"error": "Upload both PDFs"}), 400

    pdf1 = request.files["file1"]
    pdf2 = request.files["file2"]

    doc1 = fitz.open(stream=pdf1.read(), filetype="pdf")
    pdf2.stream.seek(0)
    doc2 = fitz.open(stream=pdf2.read(), filetype="pdf")

    words1 = []
    words2 = []

    # Extract words
    for page_num, page in enumerate(doc1, start=1):
        for w in page.get_text("words"):
            words1.append({"page": page_num, "word": w[4], "bbox": w[:4]})

    for page_num, page in enumerate(doc2, start=1):
        for w in page.get_text("words"):
            words2.append({"page": page_num, "word": w[4], "bbox": w[:4]})

    mismatches = []

    # Group by page and compare word sets
    pages = max(
        max(w["page"] for w in words1) if words1 else 0,
        max(w["page"] for w in words2) if words2 else 0,
    )
    for page in range(1, pages + 1):
        w1_page = [w for w in words1 if w["page"] == page]
        w2_page = [w for w in words2 if w["page"] == page]

        words1_set = {w["word"] for w in w1_page}
        words2_set = {w["word"] for w in w2_page}

        # Words unique to each PDF
        only_in_1 = [w for w in w1_page if w["word"] not in words2_set]
        only_in_2 = [w for w in w2_page if w["word"] not in words1_set]

        for w in only_in_1:
            mismatches.append({"pdf1": w})
        for w in only_in_2:
            mismatches.append({"pdf2": w})

    # Highlight mismatches in both PDFs
    highlight_mismatches(doc1, mismatches, "pdf1")
    highlight_mismatches(doc2, mismatches, "pdf2")

    # Save to memory
    buf1 = io.BytesIO()
    buf2 = io.BytesIO()
    doc1.save(buf1)
    doc2.save(buf2)
    buf1.seek(0)
    buf2.seek(0)

    return {
        "pdf1": buf1.getvalue().decode("latin1"),
        "pdf2": buf2.getvalue().decode("latin1"),
    }

@app.route("/")
def index():
    return send_from_directory(app.static_folder, "index.html")


# ------------------ START SERVER ------------------
#if __name__ == "__main__": 
 #   app.run(debug=True)